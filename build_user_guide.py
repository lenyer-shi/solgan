# -*- coding: utf-8 -*-
"""生成《水质三维荧光光谱指纹工具 - 使用指导》Word 文档."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor


def set_cell_shading(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_default_font(doc, name="Microsoft YaHei", size_pt=10.5):
    style = doc.styles["Normal"]
    style.font.name = name
    style.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size_pt)


def add_h(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Microsoft YaHei"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if level == 0:
            run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
        elif level == 1:
            run.font.color.rgb = RGBColor(0x4C, 0x72, 0xB0)
        else:
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return h


def add_p(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.name = "Microsoft YaHei"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_code_block(doc, text):
    """代码块: 等宽字体 + 浅灰背景."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9.5)

    # 浅灰背景
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F4F4F4")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_shading(hdr_cells[i], "4C72B0")
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.name = "Microsoft YaHei"
                r.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Microsoft YaHei"
                    r.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    r.font.size = Pt(9.5)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    return table


# =============================================================================
doc = Document()
set_default_font(doc)

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# ----- 封面/标题 -----
title = doc.add_heading("水质三维荧光光谱指纹工具", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in title.runs:
    r.font.name = "Microsoft YaHei"
    r.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

sub = doc.add_paragraph("使用指导  v1.0")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for r in sub.runs:
    r.font.name = "Microsoft YaHei"
    r.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_paragraph()

# ----- 1. 工具简介 -----
add_h(doc, "1. 工具简介", 1)
add_p(doc,
      "本工具是基于 PyQt5 的水质三维荧光光谱（3D-EEM, Three-Dimensional "
      "Excitation-Emission Matrix）指纹特征提取与比对系统，参考姜赞成等"
      "《基于三维荧光光谱的水体识别及组分快速解析算法研究》（光学精密工程, "
      "2025, 33(10): 1627-1637）算法实现。"
      "工具集成了完整的 EEM 预处理流水线、指纹特征提取、SQLite 指纹库管理、"
      "图形化检索与双样本比对功能，适用于工业废水、生活污水、地表水、雨水等"
      "水样的三维荧光指纹库建设与现场水质识别。")

add_h(doc, "1.1 主要功能", 2)
add_bullet(doc, "EEM 数据加载（CSV 格式，第 1 行=发射波长, 第 1 列=激发波长, 末行=DARK 背景）")
add_bullet(doc, "完整预处理流水线：DARK 扣减 → 饱和值掩蔽 → 同 Ex 通道稳健聚合 → 空白扣减 → 瑞利/拉曼散射屏蔽 → Delaunay 内插 → 二维 Savitzky-Golay 平滑")
add_bullet(doc, "指纹特征提取：FRI 五区域积分、主荧光峰位、统计描述子、标准化特征向量")
add_bullet(doc, "可视化：等高线图（可切换填充着色 / 仅等值线）、伪彩图、原始 EEM、FRI 占比柱状图")
add_bullet(doc, "指纹库管理：SQLite 持久化，支持创建 / 入库 / 删除 / 查看 / 检索")
add_bullet(doc, "相似度算法：余弦相似度 + 归一化欧式距离（权重可调）")
add_bullet(doc, "双样本一对一比对，结果以指标 + 并排等高线图呈现")

# ----- 2. 安装与启动 -----
add_h(doc, "2. 安装与启动", 1)
add_h(doc, "2.1 系统要求", 2)
add_bullet(doc, "操作系统：Windows 10/11 / macOS / Linux")
add_bullet(doc, "Python：3.8 及以上版本")
add_bullet(doc, "依赖库：numpy, pandas, matplotlib, PyQt5")

add_h(doc, "2.2 安装依赖", 2)
add_p(doc, "在 D:\\spectrum_NEW 目录下打开命令行，执行：")
add_code_block(doc, "pip install -r requirements.txt")

add_h(doc, "2.3 启动主程序", 2)
add_code_block(doc, "cd D:\\spectrum_NEW\npython eem_gui.py")
add_p(doc, "首次运行若出现旧字节码缓存问题，可先删除 __pycache__ 目录：")
add_code_block(doc, "rd /s /q __pycache__")

# ----- 3. 数据文件格式 -----
add_h(doc, "3. EEM 数据文件格式", 1)
add_p(doc,
      "工具读取约定格式的 CSV 文件，结构如下：")
add_table(doc,
    headers=["位置", "内容", "说明"],
    rows=[
        ["第 1 行第 1 列", "Ex/Em（标识列）", "可任意填写，工具忽略"],
        ["第 1 行第 2 列起", "发射波长 Em (nm)", "升序，例如 197.25, 197.57, …, 806.6"],
        ["第 2 行起，第 1 列", "激发波长 Ex (nm)", "可重复（同 Ex 多通道），如 255, 255, 265, 265, …"],
        ["其余单元格", "荧光强度计数", "16-bit ADC，0–65535（65535 视为饱和）"],
        ["末行", "DARK", "第 1 列为 'DARK'，其余为背景光谱"],
    ],
    col_widths=[3.5, 4.5, 8.5])
add_p(doc,
      "重要约定：同一 Ex 出现多次代表设备做了重复测量，因水质波动两次测量"
      "的强度可能不同。工具默认采用稳健中位数对其聚合以抑制波动影响。")

# ----- 4. 主界面布局 -----
add_h(doc, "4. 主界面布局", 1)
add_p(doc,
      "主窗口由 4 个 Tab 组成，分别对应工具的 4 类功能：")
add_table(doc,
    headers=["Tab", "名称", "用途"],
    rows=[
        ["1", "数据加载与预处理", "选择 CSV、配置参数、运行预处理与特征提取，查看可视化结果"],
        ["2", "指纹库管理", "新建/打开 SQLite 库，将当前样本入库，浏览/删除已有指纹"],
        ["3", "指纹检索", "用当前样本作为查询，按余弦+欧式综合评分排序检索库内最相似指纹"],
        ["4", "双样本比对", "任选两组 blank+sample，输出相似度并并排显示 EEM"],
    ],
    col_widths=[1.5, 4.0, 11.0])

# ----- 5. 详细操作流程 -----
add_h(doc, "5. 详细操作流程", 1)

add_h(doc, "5.1 加载 EEM 文件", 2)
add_p(doc, "切换到 Tab [1. 数据加载与预处理]，左侧 “EEM 文件加载” 区域：")
add_bullet(doc, "点击 “📂 加载空白 EEM 文件…” 选择纯水参考 CSV（如 blank.csv）")
add_bullet(doc, "点击 “📂 加载样本 EEM 文件…” 选择待测 CSV（如 industrial.csv）")
add_bullet(doc, "✕ 按钮可清除已选择的文件；标签提示当前所选文件名，悬停可看到完整路径")
add_p(doc, "也可使用菜单 “文件” → “加载空白/样本 EEM…”，对应快捷键 Ctrl+B / Ctrl+L。")

add_h(doc, "5.2 配置预处理参数", 2)
add_p(doc, "在 “预处理参数” 分组中按需调整（详细说明见第 6 章）。")

add_h(doc, "5.3 运行处理", 2)
add_p(doc, "点击蓝色 “▶ 运行预处理 + 特征提取” 按钮（或按 F5）。完成后右侧出现：")
add_bullet(doc, "等高线 (corrected)：预处理后的 EEM 等高线图，右下角可切换“填充着色 / 仅等值线”")
add_bullet(doc, "伪彩图 (corrected)：以 viridis 配色绘制的 pcolormesh 伪彩图")
add_bullet(doc, "原始 (raw)：DARK 扣减、通道聚合后但未做空白扣减/平滑的对照图")
add_bullet(doc, "FRI 占比：5 区域荧光积分占比柱状图")
add_p(doc, "下方三栏表格分别显示 FRI 五区域、主荧光峰、统计描述子的详细数值。")

add_h(doc, "5.4 入库", 2)
add_p(doc, "切换到 Tab [2. 指纹库管理]：")
add_bullet(doc, "点击 “打开/新建 .sqlite” 选择库文件位置")
add_bullet(doc, "点击 “+ 入库 (当前样本)”，依次填入指纹名称、类别（如“工业废水”）、备注")
add_bullet(doc, "表格自动刷新，显示新加入的指纹条目")

add_h(doc, "5.5 检索", 2)
add_p(doc, "切换到 Tab [3. 指纹检索]：")
add_bullet(doc, "Top-K：返回最相似的前 K 条指纹（默认 10）")
add_bullet(doc, "余弦权重 / 欧式权重：综合评分加权系数（默认 0.6 / 0.4，相加可不必为 1）")
add_bullet(doc, "点击 “🔍 用当前样本检索”，结果按综合评分降序展示")
add_bullet(doc, "选中某行，下方显示查询样本与该候选指纹的并排等高线图")

add_h(doc, "5.6 双样本比对", 2)
add_p(doc, "切换到 Tab [4. 双样本比对]：")
add_bullet(doc, "为样本 A、B 各自选择 blank 与 sample CSV")
add_bullet(doc, "点击 “▶ 计算两样本相似度” 输出余弦相似度、归一化欧式距离与综合评分")
add_bullet(doc, "下方并排显示两样本预处理后的 EEM 等高线图")

# ----- 6. 预处理参数详细说明（核心章节） -----
add_h(doc, "6. 预处理参数详细说明", 1)
add_p(doc,
      "预处理是 EEM 指纹分析的基石。原始 EEM 包含 DARK 背景、ADC 饱和、瑞利/"
      "拉曼散射、随机噪声以及水质波动等多种干扰，预处理流水线的目标是把这些"
      "干扰一一消除或抑制，得到能够直接表征水中有机组分的“干净”荧光矩阵。"
      "本章逐项说明每个可调参数的含义、推荐取值与调整策略。")

# 6.1 通道聚合
add_h(doc, "6.1 通道聚合 (agg_method)", 2)
add_p(doc, "默认值：median", bold=True)
add_p(doc,
      "作用：抑制水质波动。设备在某些 Ex（如 255 nm、265 nm）做了重复测量；"
      "因水质流动、温度漂移、瞬时颗粒物等原因，重复通道在同一 (Ex, Em) "
      "位置的强度可能差异较大。本步骤按逐 Em 维度对同 Ex 通道做稳健聚合，"
      "把它们融合为一行，使后续算法看到的是“去除瞬态波动后的代表性强度”。")
add_table(doc,
    headers=["选项", "数学含义", "适用场景"],
    rows=[
        ["median", "逐 Em 取中位数", "默认，最稳健；对个别尖峰、电气脉冲、瞬时气泡极端值不敏感"],
        ["hampel", "中位数 + MAD 离群剔除后均值", "数据有明显离群但主分布较窄时，可在保留信号细节同时过滤孤立异常"],
        ["tmean", "排序后剔除首尾 25% 的截尾均值", "样本数较多（>4 通道）且离群分布对称时使用"],
        ["mean", "简单算术平均", "实验室高精度复测、噪声主要为高斯且无离群；速度最快"],
    ],
    col_widths=[2.5, 5.5, 9.5])
add_p(doc, "调整建议：现场监测建议保持 median；若发现指纹评分异常波动，可尝试 hampel。")

# 6.2 Em 范围
add_h(doc, "6.2 Em 最小 / Em 最大 (em_range)", 2)
add_p(doc, "默认值：250 ~ 600 nm", bold=True)
add_p(doc,
      "作用：截取分析所用的发射波长区间。原始光谱通常覆盖 197–806 nm，但："
      "短波端 (<250 nm) 主要是激发光散射与仪器电子噪声；长波端 (>600 nm) "
      "几乎不含典型有机污染物的荧光信号，且容易引入二阶光栅杂散光。把分析"
      "范围收窄到 250–600 nm 既能去掉无关区域，又能加快后续计算并避免散射"
      "屏蔽时把有效信号误屏蔽。")
add_p(doc, "调整建议：")
add_bullet(doc, "若关注短波段类酪氨酸（Ex 220–250 nm）信号，且仪器支持，可把 Em 最小调到 220 nm")
add_bullet(doc, "若样本含特殊色素（如藻蓝蛋白）需要看到 650–680 nm 的发射，可把 Em 最大放宽至 700 nm")

# 6.3 Ex 上限
add_h(doc, "6.3 Ex 上限 (ex_keep_below)", 2)
add_p(doc, "默认值：600 nm", bold=True)
add_p(doc,
      "作用：丢弃过长的激发波长通道。水中有机污染物的有效荧光激发主要"
      "集中在 220–400 nm，少量物质（如某些藻类色素）可延伸到 500 nm。"
      "Ex > 600 nm 的通道（如本设备的 660、700 nm）通常是仪器为多功能"
      "保留的“非荧光分析”通道，对有机指纹不构成实质贡献，反而会引入"
      "近红外杂散光。剔除这些通道能让特征向量更聚焦于真正有机化学意义的区域。")
add_p(doc, "调整建议：常规水质监测保持 600；若需要识别叶绿素或浮游色素，调到 700–750。")

# 6.4 瑞利散射
add_h(doc, "6.4 瑞利屏蔽带宽 (rayleigh_band)", 2)
add_p(doc, "默认值：15 nm", bold=True)
add_p(doc,
      "作用：屏蔽 1 阶瑞利散射 (Em ≈ Ex) 和 2 阶瑞利散射 (Em ≈ 2·Ex)。"
      "瑞利散射强度可比有机荧光强 2–3 个数量级，会形成沿对角线的强亮带，"
      "若不去除会主导 PCA、相似度等所有下游计算。带宽指“以散射中心为对称"
      "中心向两侧屏蔽的发射波长范围”，例如 15 nm 表示 Em∈[Ex−15, Ex+15] "
      "整段都被置为 NaN，再由 Delaunay 内插补全。")
add_p(doc, "调整建议：")
add_bullet(doc, "带宽过小（<8）会有残留散射拖尾，污染下游评分")
add_bullet(doc, "带宽过大（>25）会把靠近散射边缘的真实荧光峰也覆盖掉")
add_bullet(doc, "默认 15 适用于本设备 1 nm 缝宽；缝宽更大（如 5 nm）时建议加到 20")

# 6.5 拉曼散射
add_h(doc, "6.5 拉曼屏蔽带宽 (raman_band)", 2)
add_p(doc, "默认值：15 nm", bold=True)
add_p(doc,
      "作用：屏蔽水分子的拉曼散射峰。水的 OH 伸缩振动会在某个固定的波数"
      "位移（≈ 3400 cm⁻¹）产生拉曼散射，对应发射波长由 1/Em = 1/Ex − "
      "3.4×10⁻⁴ 给出。例如 Ex = 350 nm 时拉曼峰约在 Em = 396 nm，正好"
      "落在类色氨酸/类富里酸的关键区域，必须屏蔽。本工具按上述公式自动"
      "计算每个 Ex 对应的拉曼中心，再以指定带宽屏蔽。")
add_p(doc, "调整建议：保持与瑞利同量级；离子强度极高的工业废水可减小到 10。")

# 6.6 SG 窗口
add_h(doc, "6.6 SG 窗口 (sg_window)", 2)
add_p(doc, "默认值：11 个采样点", bold=True)
add_p(doc,
      "作用：Savitzky-Golay 二维平滑的滑动窗长度，必须为奇数。窗口越大"
      "平滑越强、噪声抑制越彻底，但同时也会模糊真实荧光峰的位置和宽度。"
      "本工具依次沿 Em 方向、Ex 方向各做一次 1D SG，构成扩展二维 SG。"
      "在 Ex 通道很少（如 10 个）时，Ex 方向窗口会自动被截断到不超过通道数。")
add_p(doc, "调整建议：")
add_bullet(doc, "样本干净（信噪比高）：可下调至 7 以保留更多细节")
add_bullet(doc, "样本嘈杂（基线毛刺、强度低）：可上调至 15–21 以增强平滑")
add_bullet(doc, "改变后必须保证 sg_window > sg_poly")

# 6.7 SG 多项式阶
add_h(doc, "6.7 SG 多项式阶 (sg_poly)", 2)
add_p(doc, "默认值：3", bold=True)
add_p(doc,
      "作用：SG 滤波的局部多项式拟合阶数。阶数越高对峰形保留越好但平滑"
      "强度越弱；阶数越低平滑越强但易把窄峰变扁。3 阶能在抑制随机噪声"
      "和保留荧光峰形之间取得最优折中。")
add_p(doc, "调整建议：通常保持 3；若 sg_window 较小（如 5–7）可保持 2–3，"
      "若 sg_window 较大（>21）可上调到 4 以避免过度平滑。")

# 参数综合速查表
add_h(doc, "6.8 参数速查表", 2)
add_table(doc,
    headers=["参数", "默认", "建议范围", "影响方向"],
    rows=[
        ["agg_method", "median", "median / hampel", "↑ 抗水质波动鲁棒性"],
        ["em_range", "(250, 600)", "(220, 700)", "↑ 范围，包含更多波段；↓ 范围，加速计算"],
        ["ex_keep_below", "600 nm", "500 ~ 750", "↑ 保留更多通道；↓ 排除杂散通道"],
        ["rayleigh_band", "15 nm", "10 ~ 25", "↑ 屏蔽更彻底；↓ 保留更多真实信号"],
        ["raman_band", "15 nm", "10 ~ 20", "同上"],
        ["sg_window", "11", "7 ~ 21（奇数）", "↑ 平滑更强；↓ 保留细节"],
        ["sg_poly", "3", "2 ~ 4", "↑ 保形性更好；↓ 平滑更彻底"],
    ],
    col_widths=[3.0, 2.2, 3.2, 8.0])

# ----- 7. 算法流程图 -----
add_h(doc, "7. 算法流水线（按执行顺序）", 1)
flow = [
    "(1) 读取 CSV → BOM 剥离、首列分离 Ex 标签",
    "(2) DARK 行扣减：每个通道 - DARK 行（消除暗电流和环境光）",
    "(3) 饱和值掩蔽：≥65535 的像素置 NaN（避免污染统计量）",
    "(4) 同 Ex 通道稳健聚合：median/hampel/tmean/mean，抑制水质波动",
    "(5) Em 范围裁剪 + Ex 长波丢弃：聚焦有机分析区段",
    "(6) 空白扣减：sample_aggregated − blank_aggregated",
    "(7) 瑞利+拉曼+反斯托克斯散射区域屏蔽（→ NaN）",
    "(8) Delaunay 三角形内插（沿 Em → Ex 双向 1D 退化）补全 NaN",
    "(9) 二维 Savitzky-Golay 平滑（Em → Ex 各一次）",
    "(10) 负值截断（≥0）",
    "(11) 特征提取：FRI 五区积分 / 峰检测 / 统计 / 标准化展平向量",
]
for line in flow:
    add_bullet(doc, line)

# ----- 8. 指纹库与相似度算法 -----
add_h(doc, "8. 指纹库与相似度算法", 1)
add_h(doc, "8.1 数据库结构", 2)
add_p(doc, "SQLite 单文件数据库，主表 fingerprints 字段如下：")
add_table(doc,
    headers=["字段", "类型", "说明"],
    rows=[
        ["id", "INTEGER PK", "自增主键"],
        ["name", "TEXT UNIQUE", "指纹名（用户起）"],
        ["category", "TEXT", "类别，如“工业废水”"],
        ["source_file", "TEXT", "原始 CSV 文件名"],
        ["blank_file", "TEXT", "对应空白文件名"],
        ["created_at", "TEXT", "入库时间戳"],
        ["note", "TEXT", "备注"],
        ["ex_json / em_json", "TEXT", "Ex/Em 网格 JSON"],
        ["feature_blob", "BLOB", "float32 特征向量"],
        ["fri_json", "TEXT", "FRI 五区比例"],
        ["peaks_json", "TEXT", "主峰列表"],
        ["stats_json", "TEXT", "统计描述子"],
    ],
    col_widths=[4.0, 3.0, 9.5])

add_h(doc, "8.2 相似度公式", 2)
add_code_block(doc,
    "score = w_cos · cosine_similarity\n"
    "      + w_eu  · (1 − euclidean_distance / 2)\n\n"
    "cosine_similarity   = a · b / (‖a‖·‖b‖)\n"
    "euclidean_distance  = ‖a/‖a‖ − b/‖b‖‖    ∈ [0, 2]\n"
    "默认权重: w_cos = 0.6, w_eu = 0.4")
add_p(doc,
      "余弦相似度对幅值不敏感，刻画的是“两条指纹方向是否一致”；归一化欧式"
      "距离同样消除幅值差异，但对每个维度的差额敏感，能识别局部结构差异。"
      "二者加权可兼顾全局形状与局部细节，权重可以在检索 Tab 中实时调整。")

# ----- 9. 常见问题 -----
add_h(doc, "9. 常见问题（FAQ）", 1)
add_table(doc,
    headers=["问题", "原因 / 解决"],
    rows=[
        ["启动时报 'no module named PyQt5'",
         "未安装 GUI 依赖。运行 pip install PyQt5"],
        ["运行时报 plot_pcolor 不存在",
         "通常是缓存的 .pyc 文件未刷新。删除 __pycache__ 后重新启动"],
        ["检索结果全为 0 或异常高",
         "确认查询样本与库内指纹使用了同一台仪器、相同 Em/Ex 网格；网格不一致时工具会按最短长度对齐"],
        ["FRI 区域 I/II/III 全为 0",
         "本设备最低 Ex = 255 nm，无法激发到 220–250 nm 区段，属正常仪器限制"],
        ["等高线图过亮被对角线散射主导",
         "调小瑞利/拉曼带宽过小或未启用空白扣减；先确认 blank.csv 已加载"],
        ["重复测同一样本指纹评分波动大",
         "切换聚合方法到 hampel，或检查仪器是否在升温稳定期"],
    ],
    col_widths=[5.5, 11.0])

# ----- 10. 文件清单 -----
add_h(doc, "10. 工程文件清单", 1)
add_table(doc,
    headers=["文件", "作用"],
    rows=[
        ["eem_gui.py", "PyQt5 主程序入口"],
        ["eem_fingerprint.py", "核心算法：预处理流水线、特征提取、相似度"],
        ["fingerprint_db.py", "SQLite 指纹库管理"],
        ["eem_visualizer.py", "matplotlib 嵌入 PyQt5 的可视化组件（含 ContourPanel）"],
        ["requirements.txt", "依赖清单"],
        ["blank.csv", "纯水空白 EEM 示例"],
        ["industrial.csv", "工业废水样本 EEM 示例"],
    ],
    col_widths=[5.0, 11.5])

# ----- 11. 命令行模式 -----
add_h(doc, "11. 命令行模式（无 GUI）", 1)
add_p(doc, "对于自动化部署或批量处理，可直接调用 CLI：")
add_code_block(doc,
    "python eem_fingerprint.py --blank blank.csv --sample industrial.csv ^\n"
    "       --outdir ./report ^\n"
    "       [--agg median] [--em-min 250] [--em-max 600] [--no-plot]")
add_p(doc, "输出包括：report.txt（文字报告）、fingerprint.json（机器可读特征）、"
      "sample_eem_corrected.csv（校正矩阵）以及 4 张 PNG 图。")

# ----- 页脚信息 -----
doc.add_paragraph()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
end_run = end_p.add_run("— 文档结束 —")
end_run.font.name = "Microsoft YaHei"
end_run.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
end_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
end_run.font.size = Pt(10)

doc.save("/sessions/fervent-great-einstein/mnt/spectrum_NEW/水质三维荧光光谱指纹工具_使用指导.docx")
print("OK")
